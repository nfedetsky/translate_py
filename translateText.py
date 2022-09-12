import cv2
import pytesseract
import easyocr
import matplotlib.pyplot as plt
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

img_file = 'D:/Repository/Python/translate_py/files/test.jpg'

img = Image.open(img_file)
pytesseract.pytesseract.tesseract_cmd = r'C://Program Files/Tesseract-OCR/tesseract.exe'

im_text = pytesseract.image_to_string(img, lang='rus')

document = Document()
document = Document('D:/Repository/Python/translate_py/files/tt_1.docx')

style = document.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)
par_1 = document.add_paragraph('Распознанный текст' + '\n')
par_1.add_run(im_text)
document.save('D:/Repository/Python/translate_py/files/tt_1.docx')
